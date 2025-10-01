from io import BytesIO
from typing import List

from docx import Document


def create_word_file(title: str, paragraphs: List[str]) -> BytesIO:
    doc = Document()
    doc.add_heading(title, level=1)

    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
